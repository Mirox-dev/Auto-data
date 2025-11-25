import os
import hashlib
import zlib
from datetime import datetime
from docx import Document
from docx.shared import Pt

# Функции хэшей
def get_md5(file_path):
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        chunk = f.read(8192)
        while chunk:
            hash_md5.update(chunk)
            chunk = f.read(8192)
    return hash_md5.hexdigest()

def get_crc32(file_path):
    crc = 0
    with open(file_path, "rb") as f:
        chunk = f.read(8192)
        while chunk:
            crc = zlib.crc32(chunk, crc)
            chunk = f.read(8192)
    return f"{crc & 0xFFFFFFFF:08x}"

# Размер в байтах
def format_file_size(size_bytes):
    return f"{size_bytes} байт"

# Формат даты
def format_modification_date(file_path):
    timestamp = os.path.getmtime(file_path)
    dt = datetime.fromtimestamp(timestamp)
    months = ["января", "февраля", "марта", "апреля", "мая", "июня",
              "июля", "августа", "сентября", "октября", "ноября", "декабря"]
    return f"{dt.day} {months[dt.month-1]} {dt.year} года, {dt.hour:02d}:{dt.minute:02d}:{dt.second:02d}"

# Создание отчета
def create_word_report(directory_path, hash_type="md5"):
    folder_name = os.path.basename(os.path.normpath(directory_path))
    doc = Document()
    doc.add_heading(f"Отчет о файлах в папке: {folder_name}", level=1)

    # Таблица
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = ["Примечание (наименование файла документа)",
               "Дата и время последнего изменения файла",
               "Размер файла (байт)",
               "Значение контрольной суммы"]

    for i, header in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(header)
        run.bold = True
        hdr_cells[i].width = Pt(200)

    # Заполнение данных
    for filename in os.listdir(directory_path):
        full_path = os.path.join(directory_path, filename)
        if os.path.isfile(full_path):
            try:
                mod_date = format_modification_date(full_path)
                size = format_file_size(os.path.getsize(full_path))
                if hash_type.lower() == "md5":
                    checksum = get_md5(full_path)
                elif hash_type.lower() == "crc32":
                    checksum = get_crc32(full_path)
                else:
                    checksum = "N/A"

                row_cells = table.add_row().cells
                row_cells[0].text = filename
                row_cells[1].text = mod_date
                row_cells[2].text = size
                row_cells[3].text = checksum
            except Exception as e:
                row_cells = table.add_row().cells
                row_cells[0].text = filename
                row_cells[1].text = f"Ошибка: {e}"
                row_cells[2].text = "-"
                row_cells[3].text = "-"

    output_path = os.path.join(directory_path, "report.docx")
    doc.save(output_path)
    print(f"Отчет создан: {output_path}")

if __name__ == "__main__":
    folder = input("Введите путь к папке: ").strip('"')
    hash_choice = input("Выберите тип контрольной суммы (md5/crc32): ").strip().lower()
    if hash_choice not in ["md5", "crc32"]:
        print("Неверный выбор, используется md5 по умолчанию.")
        hash_choice = "md5"
    create_word_report(folder, hash_choice)
