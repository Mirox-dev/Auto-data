# -*- coding: utf-8 -*-
import os
import hashlib
import zlib
from datetime import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

# ---------- Хелперы ----------

def get_md5(path):
    h = hashlib.md5()
    with open(path, "rb") as f:
        chunk = f.read(8192)
        while chunk:
            h.update(chunk)
            chunk = f.read(8192)
    return h.hexdigest()

def get_crc32(path):
    crc = 0
    with open(path, "rb") as f:
        chunk = f.read(8192)
        while chunk:
            crc = zlib.crc32(chunk, crc)
            chunk = f.read(8192)
    return f"{crc & 0xFFFFFFFF:08x}"

def format_dt_with_time(ts):
    dt = datetime.fromtimestamp(ts)
    months = ["января","февраля","марта","апреля","мая","июня",
              "июля","августа","сентября","октября","ноября","декабря"]
    return "{:d} {} {:d} года, {:02d}:{:02d}:{:02d}".format(
        dt.day, months[dt.month-1], dt.year, dt.hour, dt.minute, dt.second
    )

def format_date_only_from_ts(ts):
    dt = datetime.fromtimestamp(ts)
    months = ["января","февраля","марта","апреля","мая","июня",
              "июля","августа","сентября","октября","ноября","декабря"]
    return "{:d} {} {:d}".format(dt.day, months[dt.month-1], dt.year)

def set_cell_centered(cell, text, bold=False):
    """Вставляет текст в ячейку, делает Times New Roman 12, центрирует."""
    # очищаем параграф
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    try:
        run.font.name = "Times New Roman"
        # устанавливаем также восточно-азиатский шрифт (чтобы корректно отображалось в docx)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    except Exception:
        pass
    run.font.size = Pt(12)
    run.font.bold = bold
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# ---------- Функция сборки документа по шаблону ----------

def create_gost_doc(target_paths, out_dir, out_basename, hash_type, workers_count):
    """
    target_paths: list of file paths to include (если выбран один файл - список из одного)
    out_dir: куда сохранять
    out_basename: имя файла без расширения
    hash_type: 'md5' или 'crc32'
    workers_count: целое >0
    """
    # масштаб, чтобы таблица поместилась в рабочую ширину A4 (учтены поля по 2 см)
    global ts
    SCALE = 0.95

    # ширины для пятой строки (эти ширины требуются)
    w1, w2, w3, w4 = 4.92 * SCALE, 3.79 * SCALE, 4.25 * SCALE, 3.71 * SCALE  # в см

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Заготовка таблицы: вычислим число строк
    n_files = len(target_paths)
    # структура строк:
    # 1 - header "Наименование объекта" (merge 2..4)
    # 2 - header row (4 cells)
    # 3 - same as 2 (with second cell = trimmed name)
    # 4 - like 1 with first cell = hash_type
    # 5 - column headers (4 cells)
    # 5 + n_files - file rows
    # then header for signatures (1)
    # then m worker rows
    # then last row (inform sheet)
    total_rows = 1 + 1 + 1 + 1 + 1 + n_files + 1 + workers_count + 1

    table = doc.add_table(rows=total_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Устанавливаем базовые ширины колонок согласно пятой строке (т.е. для данных по файлам)
    for i, width_cm in enumerate([w1, w2, w3, w4]):
        for cell in table.columns[i].cells:
            # Иногда docx игнорирует cell.width, но это помогает
            cell.width = Cm(width_cm)

    # helper: set cell by row index/col index
    def set_cell(r, c, text, bold=False):
        set_cell_centered(table.rows[r].cells[c], text, bold=bold)

    row_idx = 0

    # 1) Первая строка: "Наименование объекта" жирными, первая ячейка шириной 4.92 см.
    # реализуем merge второй-четвёртой ячейки
    cell0 = table.cell(row_idx, 0)
    merged_1_3 = table.cell(row_idx, 1).merge(table.cell(row_idx, 3))
    set_cell_centered(cell0, "Наименование объекта", bold=True)
    set_cell_centered(merged_1_3, "", bold=False)
    row_idx += 1

    # 2) Вторая строка: заголовки (Номер п/п, Обозначение документа, Наименование документа, Номер последнего изменения (версии))
    set_cell(row_idx, 0, "Номер п/п", bold=True)
    set_cell(row_idx, 1, "Обозначение документа", bold=True)
    set_cell(row_idx, 2, "Наименование документа", bold=True)
    set_cell(row_idx, 3, "Номер последнего изменения (версии)", bold=True)
    row_idx += 1

    # 3) Третья строка: аналогичная, пустые кроме второй ячейки = имя пользователя без 3 последних символов
    trimmed = out_basename[:-3] if len(out_basename) > 3 else out_basename
    set_cell(row_idx, 0, "")
    set_cell(row_idx, 1, trimmed)
    set_cell(row_idx, 2, "")
    set_cell(row_idx, 3, "")
    row_idx += 1

    # 4) Четвёртая строка: как первая: первая ячейка = тип шифрования, вторая пустая (merge)
    cell0 = table.cell(row_idx, 0)
    merged_1_3 = table.cell(row_idx, 1).merge(table.cell(row_idx, 3))
    set_cell_centered(cell0, hash_type.upper(), bold=False)
    set_cell_centered(merged_1_3, "", bold=False)
    row_idx += 1

    # 5) Пятая строка: заголовки колонки данных
    set_cell(row_idx, 0, "Наименование файла", bold=True)
    set_cell(row_idx, 1, "Дата и время последнего изменения файла", bold=True)
    set_cell(row_idx, 2, "Размер файла, байт", bold=True)
    set_cell(row_idx, 3, "Значение контрольной суммы", bold=True)
    row_idx += 1

    # затем n_files строк с данными
    latest_ts = None
    for p in target_paths:
        fname = os.path.basename(p)
        try:
            ts = os.path.getmtime(p)
            date_str = format_dt_with_time(ts)
            size = os.path.getsize(p)
            if hash_type.lower() == "md5":
                h = get_md5(p)
            else:
                h = get_crc32(p)
        except Exception as e:
            date_str = "Ошибка"
            size = 0
            h = "Ошибка"

        set_cell(row_idx, 0, fname)
        set_cell(row_idx, 1, date_str)
        set_cell(row_idx, 2, str(size))
        set_cell(row_idx, 3, h)

        if latest_ts is None or (isinstance(ts, (int,float)) and ts > latest_ts):
            latest_ts = ts

        row_idx += 1

    # подписи: заголовок
    set_cell(row_idx, 0, "Характер работы", bold=False)
    set_cell(row_idx, 1, "Фамилия", bold=False)
    set_cell(row_idx, 2, "Подпись", bold=False)
    set_cell(row_idx, 3, "Дата подписания", bold=False)
    row_idx += 1

    # m строк работников — пустые, но в последней колонке пишем самую позднюю дату без времени
    last_date_str = ""
    if latest_ts is not None:
        last_date_str = format_date_only_from_ts(latest_ts)
    for i in range(workers_count):
        set_cell(row_idx, 0, "")
        set_cell(row_idx, 1, "")
        set_cell(row_idx, 2, "")
        set_cell(row_idx, 3, last_date_str)
        row_idx += 1

    # последняя строка: первая ячейка 8.71 см (будем merge 0 и 1), вторая — имя файла (без расширения), третья "Лист", четвёртая "Листов"
    # merge 0 и 1
    last_row_idx = row_idx
    big = table.cell(last_row_idx, 0).merge(table.cell(last_row_idx, 1))
    set_cell_centered(big, "Информационно-удостоверяющий лист", bold=False)
    set_cell(last_row_idx, 2, out_basename)
    set_cell(last_row_idx, 3, "")  # оставляем место для "Лист/Листов" сверху
    # дополним сверху тексты "Лист" и "Листов" в отдельной строке: по требованию разместим "Лист" и "Листов" в верхней части ячейки
    # python-docx не позволяет легко позиционировать текст на 0.5 см от верха внутри ячейки; оставим стандартное выравнивание.
    # Но вставим подписи:
    # Поставим "Лист" как маленький текст сверху в ячейке 3 и "Листов" в ячейке 4
    cell_l = table.cell(last_row_idx, 2)
    cell_r = table.cell(last_row_idx, 3)
    # вставим "Лист" и "Листов" (они будут центровыми)
    set_cell_centered(cell_l, "Лист", bold=False)
    set_cell_centered(cell_r, "Листов", bold=False)

    # Сохраняем документ
    out_path = os.path.join(out_dir, out_basename + ".docx")
    doc.save(out_path)
    return out_path

# ---------- GUI ----------

class App:
    def __init__(self, root):
        self.root = root
        root.title("Создание УЛ (ГОСТ)")
        root.resizable(False, False)

        frm = ttk.Frame(root, padding=12)
        frm.grid()

        # 1) Выбор папки или файла
        ttk.Label(frm, text="Выберите папку или файл:").grid(column=0, row=0, sticky="w")
        btn = ttk.Button(frm, text="Select folder/file", command=self.select_path)
        btn.grid(column=1, row=0, padx=6, sticky="w")
        self.path_label = ttk.Label(frm, text="(путь не выбран)", width=60)
        self.path_label.grid(column=0, row=1, columnspan=2, sticky="w", pady=(2,8))

        # 2) Выбор типа шифрования
        ttk.Label(frm, text="Выберите тип шифрования:").grid(column=0, row=2, sticky="w")
        self.hash_var = tk.StringVar(value="md5")
        combo = ttk.Combobox(frm, textvariable=self.hash_var, values=["md5","crc32"], state="readonly", width=10)
        combo.grid(column=1, row=2, sticky="w")

        # 3) Наименование файла
        ttk.Label(frm, text="Наименование файла (без расширения):").grid(column=0, row=3, sticky="w", pady=(6,0))
        self.name_entry = ttk.Entry(frm, width=40)
        self.name_entry.grid(column=1, row=3, sticky="w", pady=(6,0))

        # 4) Количество работников
        ttk.Label(frm, text="Количество работников:").grid(column=0, row=4, sticky="w", pady=(6,0))
        self.workers_entry = ttk.Entry(frm, width=10)
        self.workers_entry.grid(column=1, row=4, sticky="w", pady=(6,0))

        # Кнопка создать
        create_btn = ttk.Button(frm, text="Создать УЛ", command=self.on_create)
        create_btn.grid(column=0, row=5, pady=(12,0))

        # Статус
        self.status = ttk.Label(frm, text="", foreground="green")
        self.status.grid(column=0, row=6, columnspan=2, pady=(8,0))

        # внутренние
        self.selected_path = None
        self.selected_is_folder = False

    def select_path(self):
        # сначала выбор файла, если отмена — выбор папки
        pfile = filedialog.askopenfilename(title="Выберите файл (отмена -> выбрать папку)")
        if pfile:
            self.selected_path = pfile
            self.selected_is_folder = False
            self.path_label.config(text=pfile)
            return
        pdir = filedialog.askdirectory(title="Выберите папку")
        if pdir:
            self.selected_path = pdir
            self.selected_is_folder = True
            self.path_label.config(text=pdir)
            return
        self.selected_path = None
        self.path_label.config(text="(путь не выбран)")

    def validate(self):
        if not self.selected_path:
            return False, "Не выбран путь."
        name = self.name_entry.get().strip()
        if not name:
            return False, "Не заполнено имя файла."
        workers = self.workers_entry.get().strip()
        if not (workers.isdigit() and int(workers) > 0):
            return False, "Количество работников введено неверно."
        if self.selected_is_folder:
            if not os.path.isdir(self.selected_path):
                return False, "Указанная папка не существует."
        else:
            if not os.path.isfile(self.selected_path):
                return False, "Указанный файл не найден."
        return True, ""

    def on_create(self):
        ok, msg = self.validate()
        if not ok:
            messagebox.showerror("Ошибка", "Не все поля заполнены или заполнены неправильно.\n" + msg)
            return

        hash_type = self.hash_var.get().lower()
        out_name = self.name_entry.get().strip()
        workers = int(self.workers_entry.get().strip())

        # подготовка списка файлов
        files = []
        if self.selected_is_folder:
            for nm in sorted(os.listdir(self.selected_path)):
                fp = os.path.join(self.selected_path, nm)
                if os.path.isfile(fp):
                    files.append(fp)
            save_dir = self.selected_path
        else:
            files = [self.selected_path]
            save_dir = os.path.dirname(self.selected_path) or os.getcwd()

        if len(files) == 0:
            messagebox.showerror("Ошибка", "В выбранной папке нет файлов для обработки.")
            return

        try:
            out_path = create_gost_doc(files, save_dir, out_name, hash_type, workers)
        except Exception as e:
            messagebox.showerror("Ошибка", "Ошибка при создании файла:\n" + str(e))
            return

        # очистка полей
        self.selected_path = None
        self.selected_is_folder = False
        self.path_label.config(text="(путь не выбран)")
        self.name_entry.delete(0, tk.END)
        self.workers_entry.delete(0, tk.END)
        self.hash_var.set("md5")
        self.status.config(text="Файл {}.docx создан в папке {}".format(out_name, os.path.dirname(out_path)))
        messagebox.showinfo("Готово", "Файл {}.docx создан в папке {}".format(out_name, os.path.dirname(out_path)))

# ---------- запуск ----------

def main():
    root = tk.Tk()
    app = App(root)
    root.mainloop()

if __name__ == "__main__":
    main()
