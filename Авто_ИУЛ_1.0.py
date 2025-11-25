# -*- coding: utf-8 -*-
import os
import hashlib
import zlib
from datetime import datetime
from tkinter.ttk import Combobox, Style
from tkinter import Tk, Label, Button, Entry, StringVar, W, E, messagebox
from tkinter import filedialog
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

# -------------------- Хэш-функции и утилиты --------------------

def get_md5(file_path):
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def get_crc32(file_path):
    crc = 0
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            crc = zlib.crc32(chunk, crc)
    return f"{crc & 0xFFFFFFFF:08x}"

def format_file_size(size_bytes):
    return str(size_bytes)  # байты как число (в шаблоне сказано "Размер файла, байт")

def format_modification_date(file_path):
    timestamp = os.path.getmtime(file_path)
    dt = datetime.fromtimestamp(timestamp)
    months = ["января", "февраля", "марта", "апреля", "мая", "июня",
              "июля", "августа", "сентября", "октября", "ноября", "декабря"]
    return f"{dt.day} {months[dt.month-1]} {dt.year} года, {dt.hour:02d}:{dt.minute:02d}:{dt.second:02d}"

def format_date_no_time(file_path):
    timestamp = os.path.getmtime(file_path)
    dt = datetime.fromtimestamp(timestamp)
    months = ["января", "февраля", "марта", "апреля", "мая", "июня",
              "июля", "августа", "сентября", "октября", "ноября", "декабря"]
    return f"{dt.day} {months[dt.month-1]} {dt.year}"

# -------------------- Логика создания отчёта в Word --------------------

def create_ul_report(target_path, is_file, hash_type, output_name, workers_count):
    # Определяем папку для сохранения и список файлов для отчёта
    if is_file:
        if not os.path.isfile(target_path):
            raise FileNotFoundError("Указанный файл не найден.")
        folder = os.path.dirname(os.path.abspath(target_path))
        files = [os.path.basename(target_path)]
        files_full = [target_path]
    else:
        if not os.path.isdir(target_path):
            raise FileNotFoundError("Указанная папка не найдена.")
        folder = os.path.abspath(target_path)
        # Список только обычных файлов (не директорий), без скрытых фильтров
        files_full = [os.path.join(folder, f) for f in os.listdir(folder) if os.path.isfile(os.path.join(folder, f))]
        files = [os.path.basename(p) for p in files_full]

    # Собираем данные по каждому файлу
    file_rows = []
    latest_ts = None
    for fpath in files_full:
        try:
            name = os.path.basename(fpath)
            mod_dt_full = format_modification_date(fpath)
            size = format_file_size(os.path.getsize(fpath))
            checksum = get_md5(fpath) if hash_type.upper() == "MD5" else get_crc32(fpath)
            ts = os.path.getmtime(fpath)
            if latest_ts is None or ts > latest_ts:
                latest_ts = ts
            file_rows.append((name, mod_dt_full, size, checksum))
        except Exception as e:
            # в случае ошибки - записываем строку с пометкой об ошибке
            file_rows.append((os.path.basename(fpath), f"Ошибка: {e}", "-", "-"))

    # если нет файлов (в папке пусто) — latest_date будет None, сделаем заглушку
    if latest_ts is None:
        latest_date_str = ""
    else:
        dt = datetime.fromtimestamp(latest_ts)
        months = ["января", "февраля", "марта", "апреля", "мая", "июня",
                  "июля", "августа", "сентября", "октября", "ноября", "декабря"]
        latest_date_str = f"{dt.day} {months[dt.month-1]} {dt.year}"

    # Создаём docx
    doc = Document()
    # общий шрифт для всех вставляемых параграфов — будем назначать ран по мере добавления
    # Добавим таблицу: Кол-во строк = фиксированные строки + n файлов + (1 + m) для подписей + последняя строка
    n = len(file_rows)
    m = workers_count

    # Рассчитаем необходимое количество строк:
    # 1 первая строка (Наименование объекта)
    # 1 вторая строка (шапка: Номер п/п и т.д.)
    # 1 третья строка (пустая кроме второй ячейки с именем файла без 3 символов)
    # 1 четвёртая строка (тип шифрования)
    # 1 пятая строка (заголовок столбцов: Наименование файла, Дата..., Размер..., Значение..)
    # n строк с файлами
    # 1 строка - заголовки для подписи: "Характер работы", ...
    # m строк для работников
    # 2 финальные строки (Информационно-удостоверяющий лист)
    total_rows = 1 + 1 + 1 + 1 + 1 + n + 1 + m + 1 + 1

    table = doc.add_table(rows=total_rows, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Установим ширины столбцов согласно шаблону (в см):
    # Все колонки
    col_widths = [1.93, 2.98, 3.79, 4.25, 1.53, 2.17]  # см
    # первая и четвёртая строка: первая ячейка 4.92, вторая 11.75 (занимает места остальных колонок)
    # последняя строка: первая 8.71, вторая 4.25, третья 1.53, четвёртая 2.17
    # строки с файлами и работниками: первая 4.92, вторая 3.79, третья 4.25, четвёртая 3.71

    # Установим ширины таблицы по столбцам (учтём, что 11.75 будет логически в двух колонках):
    # Чтобы обеспечить фиксированные визуальные ширины, установим ширины столбцов как basis:
    # используем колонки для шаблона: [1.93, 2,98, 3.79, 4.25, 1.53, 2.17] как базу
    for idx, width_cm in enumerate(col_widths):
        for cell in table.columns[idx].cells:
            cell.width = Cm(width_cm)

    # Функция для оформления текста в ячейке (Times New Roman, 12, выравнивание по центру, вертикальное центрирование)
    def set_cell_text(cell_cage, text, bold=False, spacing_before_cm=0):
        cell_cage.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # очищаем существующие параграфы
        para = cell_cage.paragraphs[0]
        para.clear()  # python-docx >=0.8.11 поддерживает clear()
        run_cell = para.add_run(text)
        run_cell.font.name = 'Times New Roman'
        run_cell.font.size = Pt(12)
        run_cell.bold = bold
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if spacing_before_cm and spacing_before_cm > 0:
            # set spacing before to emulate 0.5cm vertical offset
            para.space_before = Pt(spacing_before_cm * 28.3464567)  # 1 cm ≈ 28.3464567 pt

    def merge_cells(sheet, row_index, cols):
        first = cols[0]
        for col in cols[1:]:
            sheet.rows[row_index].cells[first].merge(sheet.rows[row_index].cells[col])

    # Индекс строки в таблице (0-based)
    r = 0

    # 1) Первая строка: "Наименование объекта" (первая ячейка шириной 4.92 см), вторая ячейка 11.75 см пустая
    # Реализуем: первая колонка - текст, вторая колонка объединяем визуально две колонки (2-4) оставим пустыми,
    # но пользователь просил именно вторая ячейка шириной 11.75 — это примерно суммарная ширина колонок 2+3+4
    merge_cells(table, r, [0, 1])
    set_cell_text(table.rows[r].cells[0], "Наименование объекта", bold=True)
    # Оставляем остальные пустыми (вторая должна быть шириной 11.75 см = 3 cols sum)
    merge_cells(table, r, [2, 3, 4, 5])
    set_cell_text(table.rows[r].cells[2], "")
    r += 1

    # 2) Вторая строка: "Номер п/п" (1.93), "Обозначение документа" (2.98), "Наименование документа" (8.04), "Номер последнего изменения (версии)" (3.71)
    # Соответствие: у нас базовые колонны другие; чтобы приблизить, просто запишем тексты и оставим ширины базовых колонок
    set_cell_text(table.rows[r].cells[0], "Номер п/п", bold=True)
    set_cell_text(table.rows[r].cells[1], "Обозначение документа", bold=True)
    merge_cells(table, r, [2, 3])
    set_cell_text(table.rows[r].cells[2], "Наименование документа", bold=True)
    merge_cells(table, r, [4, 5])
    set_cell_text(table.rows[r].cells[4], "Номер последнего изменения (версии)", bold=True)
    r += 1

    # 3) Третья строка: такие же ячейки, пустые, кроме второй ячейки: во второй ячейке имя файла, введённое пользователем, но без 3 последних символов.
    name_without_last3 = output_name[:-3] if len(output_name) > 3 else output_name
    set_cell_text(table.rows[r].cells[0], "")
    set_cell_text(table.rows[r].cells[1], name_without_last3)
    merge_cells(table, r, [2, 3])
    set_cell_text(table.rows[r].cells[2], "")
    merge_cells(table, r, [4, 5])
    set_cell_text(table.rows[r].cells[4], "")
    r += 1

    # 4) Четвёртая строка: такие же, как первая. В первой ячейке тип шифрования (MD5/CRC32), вторая пустая.
    merge_cells(table, r, [0, 1])
    set_cell_text(table.rows[r].cells[0], hash_type.upper())
    merge_cells(table, r, [2, 3, 4, 5])
    set_cell_text(table.rows[r].cells[2], "")
    r += 1

    # 5) Пятая строка: заголовки столбцов с указанными ширинами: "Наименование файла", "Дата и время последнего изменения файла", "Размер файла, байт", "Значение контрольной суммы"
    merge_cells(table, r, [0, 1])
    set_cell_text(table.rows[r].cells[0], "Наименование файла", bold=True)
    set_cell_text(table.rows[r].cells[2], "Дата и время последнего изменения файла", bold=True)
    set_cell_text(table.rows[r].cells[3], "Размер файла, байт", bold=True)
    merge_cells(table, r, [4, 5])
    set_cell_text(table.rows[r].cells[4], "Значение контрольной суммы", bold=True)
    r += 1

    # Далее n строк: данные о каждом файле. Ширина совпадает с пятой строкой (мы уже задали базовые колонки)
    for name, mod_dt_full, size, checksum in file_rows:
        merge_cells(table, r, [0, 1])
        set_cell_text(table.rows[r].cells[0], name)
        set_cell_text(table.rows[r].cells[2], mod_dt_full)
        set_cell_text(table.rows[r].cells[3], size)
        merge_cells(table, r, [4, 5])
        set_cell_text(table.rows[r].cells[4], checksum)
        r += 1

    # После всех файлов: строка с заголовками для подписи
    merge_cells(table, r, [0, 1])
    set_cell_text(table.rows[r].cells[0], "Характер работы", bold=True)
    set_cell_text(table.rows[r].cells[2], "Фамилия", bold=True)
    set_cell_text(table.rows[r].cells[3], "Подпись", bold=True)
    merge_cells(table, r, [4, 5])
    set_cell_text(table.rows[r].cells[4], "Дата подписания", bold=True)
    r += 1

    # m строк для работников — в последнем столбце вставляем самую позднюю дату (без времени) для всех строк
    for i in range(m):
        merge_cells(table, r, [0, 1])
        set_cell_text(table.rows[r].cells[0], "")
        set_cell_text(table.rows[r].cells[2], "")
        set_cell_text(table.rows[r].cells[3], "")
        merge_cells(table, r, [4, 5])
        set_cell_text(table.rows[r].cells[4], latest_date_str if latest_date_str else "")
        r += 1

    # Последняя строка: первая ячейка шириной 8.71 см с надписью "Информационно-удостоверяющий лист",
    # вторая ячейка 4.25 см — вставляем имя файла, под которым будут сохранять, без расширения.
    # Третья 1.53 см и четвёртая 2.17 см; сверху в третьей "Лист", в четвёртой "Листов".
    # Эмулируем это текстовым заполнением и добавим отступ сверху 0.5 см для 3-й и 4-й ячеек, чтобы была вертикальная "граница"
    # Последние две строки объединяю вручную
    base_name_no_ext = os.path.splitext(output_name)[0]
    merge_cells(table, r, [0, 1, 2])
    r += 1
    merge_cells(table, r, [0, 1, 2])
    table.rows[r-1].cells[0].merge(table.rows[r].cells[0])
    table.rows[r-1].cells[3].merge(table.rows[r].cells[3])
    set_cell_text(table.rows[r-1].cells[0], "Информационно-удостоверяющий лист")
    set_cell_text(table.rows[r-1].cells[3], base_name_no_ext)
    # Для 3 и 4 ячеек добавим сверху spacing_before ~0.5 см (см. Заметку в начале)
    set_cell_text(table.rows[r-1].cells[4], "Лист")
    set_cell_text(table.rows[r-1].cells[5], "Листов")
    r += 1

    # Применим шрифт ко всем оставшимся параграфам (на всякий случай)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # Сохраняем файл
    output_filename = f"{output_name}.docx" if not output_name.lower().endswith(".docx") else output_name
    output_path = os.path.join(folder, output_filename)
    doc.save(output_path)
    return output_path

# -------------------- GUI (Tkinter) --------------------

class App:
    def __init__(self, root):
        self.root = root
        root.title("Создание информационно-удостоверяющего листа (УЛ)")

        # Размер окна
        root.geometry("1200x400")
        root.resizable(True, True)

        # Шрифт для всех виджетов
        self.my_font = ("Times New Roman", 18)

        # Стиль для ttk виджетов
        style = Style()
        style.configure("TButton", font=self.my_font)
        style.configure("TLabel", font=self.my_font)
        style.configure("TEntry", font=self.my_font)
        style.configure("TCombobox",
                        padding=10,  # отступ внутри виджета (px)
                        font=self.my_font)

        # Переменные
        self.selected_path = StringVar()
        self.selected_is_file = False
        self.hash_choice = StringVar(value="MD5")
        self.output_name = StringVar()
        self.workers = StringVar()

        # UI элементы
        Label(root, text="Выберите папку или файл:", font=self.my_font).grid(row=0, column=0, sticky=W, padx=10,
                                                                             pady=10)
        self.path_entry = Entry(root, textvariable=self.selected_path, width=45, font=self.my_font)
        self.path_entry.grid(row=0, column=1, columnspan=1, padx=10, pady=10, sticky=W + E)

        # Две отдельные кнопки для выбора файла и папки
        self.select_file_btn = Button(root, text="Выбрать файл", font=self.my_font, width=15, command=self.select_file)
        self.select_file_btn.grid(row=0, column=2, padx=10, pady=10)

        self.select_folder_btn = Button(root, text="Выбрать папку", font=self.my_font, width=15,
                                        command=self.select_folder)
        self.select_folder_btn.grid(row=1, column=2, padx=10, pady=10)

        Label(root, text="Выберите тип шифрования:", font=self.my_font).grid(row=1, column=0, sticky=W, padx=10,
                                                                             pady=10)
        self.hash_combo = Combobox(root, textvariable=self.hash_choice, values=["MD5", "CRC32"], state="readonly",
                                   width=20)
        self.hash_combo.grid(row=1, column=1, padx=10, pady=10, sticky=W)
        self.hash_combo.current(0)

        Label(root, text="Наименование файла (без .docx):", font=self.my_font).grid(row=2, column=0, sticky=W, padx=10,
                                                                                    pady=10)
        Entry(root, textvariable=self.output_name, width=30, font=self.my_font).grid(row=2, column=1, padx=10, pady=10)

        Label(root, text="Количество работников:", font=self.my_font).grid(row=3, column=0, sticky=W, padx=10, pady=10)
        Entry(root, textvariable=self.workers, width=10, font=self.my_font).grid(row=3, column=1, padx=10, pady=10,
                                                                                 sticky=W)

        Button(root, text="Создать УЛ", font=self.my_font, command=self.on_create).grid(row=4, column=0, columnspan=4,
                                                                                        pady=20)
    def select_file(self):
        p = filedialog.askopenfilename(title="Выберите файл")
        if p:
            self.selected_path.set(p)
            self.selected_is_file = True

    def select_folder(self):
        p = filedialog.askdirectory(title="Выберите папку")
        if p:
            self.selected_path.set(p)
            self.selected_is_file = False

    def on_create(self):
        path = self.selected_path.get().strip()
        hash_type = self.hash_choice.get().strip().upper()
        out_name = self.output_name.get().strip()
        workers_str = self.workers.get().strip()

        # Валидация полей
        if not path or not out_name or not workers_str:
            messagebox.showerror("Ошибка", "Не все поля заполнены или заполнены неправильно")
            return
        # проверка существования пути
        if self.selected_is_file:
            if not os.path.isfile(path):
                messagebox.showerror("Ошибка", "Не все поля заполнены или заполнены неправильно")
                return
        else:
            if not os.path.isdir(path):
                messagebox.showerror("Ошибка", "Не все поля заполнены или заполнены неправильно")
                return
        # работники натуральное число
        try:
            workers_count = int(workers_str)
            if workers_count <= 0:
                raise ValueError()
        except Exception:
            messagebox.showerror("Ошибка", "Количество работников должно быть натуральным числом")
            return

        # Создаём отчёт
        try:
            out_path = create_ul_report(path, self.selected_is_file, hash_type, out_name, workers_count)
        except Exception as e:
            messagebox.showerror("Ошибка при создании", f"Не удалось создать файл: {e}")
            return

        # Очистить поля и показать сообщение об успехе
        self.selected_path.set("")
        self.selected_is_file = False
        self.hash_choice.set("MD5")
        self.output_name.set("")
        self.workers.set("")
        messagebox.showinfo("Готово", f"Файл {os.path.basename(out_path)} создан в папке {os.path.dirname(out_path)}")

# -------------------- Запуск приложения --------------------

def main():
    root = Tk()
    app = App(root)
    root.mainloop()

if __name__ == "__main__":
    main()
